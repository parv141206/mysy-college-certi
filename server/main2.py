from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt
from flask_cors import CORS
from docx2pdf import convert
app = Flask(__name__)
CORS(app)
@app.route('/generate-doc', methods=['POST'])
def generate_word_doc():
    data = request.get_json()
    name = data['name']
    eno = data['eno']
    year = data['year']
    branch = data['branch']
    method = data['method']
    fee = data['fee']
    is_admitted = data['is_admitted']
    gender = data['gender']

    document = Document("input.docx")
    mis = "Mr." if gender == "male" else "Ms."
    heshe = "He" if gender == "male" else "She"
    
    placeholders = {
        "MIS": mis,
        "NAMEOFSTUDENT": name,
        "ENO" : str(eno) ,
        "YEARINPUT": str(year) ,
        "BRANCH": branch, 
        "METHOD": method, 
        "HESHE": heshe,
        "AMOUNTOFFEE": str(fee) ,
        "ADMIT": is_admitted
    }

    for para in document.paragraphs:
        for run in para.runs:
            print(run.text)
            for key, value in placeholders.items():
                if key.strip() in run.text.strip() or key in run.text:
                    run.text = run.text.replace(key, value)
                    run.bold = True

    for para in document.paragraphs:
        for run in para.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14) 
    
    document.save('output-dummy.docx')
    convert('output-dummy.docx', 'output-dummy.pdf')

    return send_file('output-dummy.pdf', as_attachment=True)
@app.route("/", methods=['GET'])
def index():
    return "Hello, World!"
if __name__ == '__main__':
    app.run(debug=False)