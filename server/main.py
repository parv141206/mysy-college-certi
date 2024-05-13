from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt
from flask_cors import CORS
from docx2pdf import convert
import pythoncom
import threading
import logging
from concurrent_log_handler import ConcurrentRotatingFileHandler

app = Flask(__name__)
CORS(app)

log_handler = ConcurrentRotatingFileHandler('app.log', maxBytes=10000, backupCount=3)
log_handler.setLevel(logging.INFO)
app.logger.addHandler(log_handler)
app.logger.setLevel(logging.INFO)

def generate_and_convert(data, number_str):
    pythoncom.CoInitialize()  # Initialize COM for Windows
    
    document = Document("input_r.docx")
    
    fee_paid_in_sem = data["in_semester"]
    serial_year = "2023-2024"
    serial_year_int = 2023
    name = data['name']
    eno = data["enrollment_number"]
    acad_year = data["academic_year"]
    p_acad_year = data["previous_academic_year"]
    branch = data["branch"]
    amount = data["amount"]
    gender = data["gender"]
    hostel = data["hosteler"]
    mis = "Mr." if gender == "male" else "Ms."
    heshe = "He" if gender == "male" else "She"
    feepaid = data["fee_paid"]
    passing_year_1 = data["passing_year_1"]
    passing_sem_1 = data["passing_sem_1"]
    passing_sem_2 = data["passing_sem_2"]
    percentile_1 = data["percentile_1"]
    spi_1 = data["spi_1"]
    spi_2 = data["spi_2"]
    attempts_1 = data["attempts_1"]
    attempts_2 = data["attempts_2"]
    placeholders = {
        "SERI": str(f"VPMP/MYSY/868/{number_str}/{serial_year}"),
        "MIS": mis,
        "NAMEOFSTUDENT": name,
        "ENO": str(eno),
        "YOFATEN": str(serial_year_int - 1),
        "ACADYEAR": str(acad_year),
        "BRANCH": branch,
        "AMOUNTOFMONEY": str(amount),
        "PRYEAR": p_acad_year,
        "HESHE": heshe,
        "HOSTEL": hostel,
        "FEEPAID": feepaid,
        "FINSEM": fee_paid_in_sem,
        "PY_O": passing_year_1,
        "PSAM_O": passing_sem_1,
        "PSAM_S": passing_sem_2,
        "PERTILE_O": percentile_1,
        "SPI_O": spi_1,
        "SPI_S": spi_2,
        "ATM_O": attempts_1,
        "ATM_S": attempts_2,
    }

    for para in document.paragraphs:
        for run in para.runs:
            for key, value in placeholders.items():
                if key.strip() in run.text.strip() or key in run.text:
                    run.text = run.text.replace(key, value)
                    run.bold = True

    for para in document.paragraphs:
        for run in para.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)

    document.save(f'files/re/{number_str}.docx')
    convert(f'files/re/{number_str}.docx', f'files/re/{number_str}.pdf')

    with open("counter_re.txt", "w") as f:
        f.write(str(int(number_str) + 1))

    pythoncom.CoUninitialize()  
    return f'files/re/{number_str}.pdf'

@app.route('/test', methods=['POST'])
def handle_request():
    data = request.get_json()
    with open("./counter_re.txt", "r") as f:
        number_str = int(f.read().strip())

    thread = threading.Thread(target=generate_and_convert, args=(data, number_str))
    thread.start()
    thread.join()

    file_path = f'files/re/{number_str}.pdf'
    return send_file(file_path, as_attachment=True)

@app.route("/", methods=['GET'])
def index():
    return "Hello, World!"
@app.route('/generate-doc', methods=['POST'])
def generate_word_doc():
    pythoncom.CoInitialize() 
    with open("./counter.txt", "r") as f:
        number_str = int(f.read().strip())
    data = request.get_json()
    name = data['name']
    eno = data['eno']
    year = data['year']
    branch = data['branch']
    method = data['method']
    fee = data['fee']
    is_admitted = data['is_admitted']
    gender = data['gender']

    document = Document("input.docx")
    mis = "Mr." if gender == "male" else "Ms."
    heshe = "He" if gender == "male" else "She"
    
    placeholders = {
        "SRNOI": str(number_str),
        "MIS": mis,
        "NAMEOFSTUDENT": name,
        "ENO" : str(eno) ,
        "YEARINPUT": str(year) ,
        "BRANCH": branch, 
        "METHOD": method, 
        "HESHE": heshe,
        "AMOUNTOFFEE": str(fee) ,
        "ADMIT": is_admitted
    }

    for para in document.paragraphs:
        for run in para.runs:
            print(run.text)
            for key, value in placeholders.items():
                if key.strip() in run.text.strip() or key in run.text:
                    run.text = run.text.replace(key, value)
                    run.bold = True

    for para in document.paragraphs:
        for run in para.runs:
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14) 
    
    document.save(f'./files/fresh/{number_str}.docx')
    convert(f'./files/fresh/{number_str}.docx', f'./files/fresh/{number_str}.pdf')
    pythoncom.CoUninitialize()
    with open("counter.txt", "w") as f:
        f.write(str(int(number_str) + 1))
    
    return send_file(f'./files/fresh/{number_str}.pdf', as_attachment=True)

@app.route("/shutdown", methods=['GET'])
def shutdown():
    func = request.environ.get('werkzeug.server.shutdown')
    if func is None:
        raise RuntimeError('Not running with the Werkzeug Server')
    func()
    return 'Server shutting down...'

if __name__ == '__main__':
    app.run(debug=True)
